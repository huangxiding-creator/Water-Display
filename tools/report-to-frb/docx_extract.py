#!/usr/bin/env python3
"""
report-to-frb — Compile a water-conservancy informatization design-report
chapter (.docx or .pdf) into a structured Functional Requirements Breakdown
(JSON) that feeds the Water-Display site generator.

Strategy:
  - .docx preferred (real text + tables + headings via style/regex)
  - .pdf fallback (text via pymupdf; no tables)
  - Headings detected by regex on numbering patterns like "15.5.3.2"
  - Tables extracted with their nearest preceding heading as context

Usage:
    python docx_extract.py <report.docx|report.pdf> [out.json]

The output is a draft FRB — a human reviews it, then maps it into a
project.config.ts. This is the "config-driven + report compiler" decision
locked in the approved proposal (Q1).
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any

# Force UTF-8 stdout so emoji / Chinese don't crash on Windows GBK consoles.
try:
    sys.stdout.reconfigure(encoding='utf-8')  # type: ignore[attr-defined]
    sys.stderr.reconfigure(encoding='utf-8')  # type: ignore[attr-defined]
except (AttributeError, ValueError):
    pass

HEADING_RE = re.compile(r'^(?P<num>\d+(?:\.\d+){0,4})\s+(?P<title>\S.*)$')


def extract_docx(path: Path) -> dict[str, Any]:
    import docx  # type: ignore
    from docx.oxml.ns import qn

    doc = docx.Document(str(path))
    headings: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []

    current_heading = ''
    para_idx = 0
    tbl_idx = 0
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            p = doc.paragraphs[para_idx]
            para_idx += 1
            txt = p.text.strip()
            m = HEADING_RE.match(txt)
            if m and len(m.group('num').split('.')) >= 2:
                current_heading = txt[:120]
                headings.append({'number': m.group('num'), 'title': m.group('title')[:120], 'raw': txt[:120]})
        elif child.tag == qn('w:tbl') and tbl_idx < len(doc.tables):
            t = doc.tables[tbl_idx]
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            tables.append({
                'index': tbl_idx,
                'context_heading': current_heading,
                'rows': len(rows),
                'cols': len(rows[0]) if rows else 0,
                'header': rows[0] if rows else [],
                'data': rows[1:] if len(rows) > 1 else [],
            })
            tbl_idx += 1

    return {
        'source': str(path),
        'format': 'docx',
        'paragraph_count': len(doc.paragraphs),
        'heading_count': len(headings),
        'table_count': len(tables),
        'headings': headings,
        'tables': tables,
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    import fitz  # type: ignore

    doc = fitz.open(str(path))
    headings: list[dict[str, Any]] = []
    pages: list[dict[str, Any]] = []
    for i, page in enumerate(doc):
        text = page.get_text('text')
        pages.append({'page': i + 1, 'text': text})
        for line in text.splitlines():
            line = line.strip()
            m = HEADING_RE.match(line)
            if m and len(m.group('num').split('.')) >= 2:
                headings.append({'number': m.group('num'), 'title': m.group('title')[:120], 'page': i + 1})
    return {
        'source': str(path),
        'format': 'pdf',
        'page_count': len(doc),
        'heading_count': len(headings),
        'table_count': 0,
        'headings': headings,
        'pages': pages,
    }


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__)
        return 1
    src = Path(sys.argv[1]).resolve()
    if not src.exists():
        print(f'❌ File not found: {src}', file=sys.stderr)
        return 1

    if src.suffix.lower() == '.docx':
        data = extract_docx(src)
    elif src.suffix.lower() == '.pdf':
        data = extract_pdf(src)
    else:
        print(f'❌ Unsupported format: {src.suffix} (use .docx or .pdf)', file=sys.stderr)
        return 1

    out = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else src.with_suffix('.frb.json')
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f'✅ Extracted FRB draft:')
    print(f'   Source : {data["source"]}')
    print(f'   Format : {data["format"]}')
    print(f'   Headings: {data["heading_count"]}')
    print(f'   Tables : {data["table_count"]}')
    print(f'   Output : {out}')
    print(f'\n   Next: review {out.name}, then map into project.config.ts')
    return 0


if __name__ == '__main__':
    sys.exit(main())
